from aimakerspace.text_utils import DocxLoader
import os

# Create a sample DOCX file for testing
sample_docx_path = "tests/sample_test.docx"

# Only create the DOCX if it doesn't exist
if not os.path.exists(sample_docx_path):
    try:
        import docx
        doc = docx.Document()
        doc.add_heading('Test DOCX', 0)
        doc.add_paragraph('This is a test DOCX file.')
        doc.save(sample_docx_path)
    except ImportError:
        print("python-docx is required to generate a sample DOCX. Please install it with 'uv pip install python-docx'.")
        exit(1)

# Test the DocxLoader
loader = DocxLoader(sample_docx_path)
loader.load()
docs = loader.documents

print("Loaded documents:")
for doc in docs:
    print(doc) 