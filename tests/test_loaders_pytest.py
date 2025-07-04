import os
import pytest
from aimakerspace.text_utils import MarkdownLoader, PDFLoader, DocxLoader, ImageLoader
import re

# For PDF and DOCX generation
import sys

@pytest.fixture
def sample_md_file(tmp_path):
    file = tmp_path / "sample.md"
    content = "# Test Markdown\n\nThis is a test file.\n- Item 1\n- Item 2"
    file.write_text(content, encoding="utf-8")
    return file, content

@pytest.fixture
def sample_pdf_file(tmp_path):
    file = tmp_path / "sample.pdf"
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
    except ImportError:
        pytest.skip("reportlab is required for PDF test")
    c = canvas.Canvas(str(file), pagesize=letter)
    c.drawString(100, 750, "Test PDF")
    c.drawString(100, 735, "This is a test PDF file.")
    c.save()
    return file

@pytest.fixture
def sample_docx_file(tmp_path):
    file = tmp_path / "sample.docx"
    try:
        import docx
    except ImportError:
        pytest.skip("python-docx is required for DOCX test")
    doc = docx.Document()
    doc.add_heading('Test DOCX', 0)
    doc.add_paragraph('This is a test DOCX file.')
    doc.save(str(file))
    return file

@pytest.fixture(params=["png", "jpg", "jpeg"])
def sample_image_file(tmp_path, request):
    ext = request.param
    file = tmp_path / f"sample.{ext}"
    try:
        from PIL import Image, ImageDraw
    except ImportError:
        pytest.skip("Pillow is required for image test")
    img = Image.new('RGB', (200, 60), color = (255, 255, 255))
    d = ImageDraw.Draw(img)
    text = f"Test {ext.upper()}"
    d.text((10, 10), text, fill=(0, 0, 0))
    img.save(str(file))
    return file, text

def normalize_text(text):
    # Lowercase, remove punctuation, and collapse whitespace
    text = text.lower()
    text = re.sub(r'[^a-z0-9\s]', '', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text

def test_markdown_loader(sample_md_file):
    file, content = sample_md_file
    loader = MarkdownLoader(str(file))
    loader.load()
    docs = loader.documents
    assert any("Test Markdown" in doc for doc in docs)
    assert any("Item 1" in doc for doc in docs)

def test_pdf_loader(sample_pdf_file):
    loader = PDFLoader(str(sample_pdf_file))
    loader.load()
    docs = loader.documents
    print("Extracted PDF text:", docs)
    normalized_docs = [normalize_text(doc) for doc in docs]
    assert any("test pdf" in doc for doc in normalized_docs)
    assert any("test pdf file" in doc for doc in normalized_docs)

def test_docx_loader(sample_docx_file):
    loader = DocxLoader(str(sample_docx_file))
    loader.load()
    docs = loader.documents
    assert any("Test DOCX" in doc for doc in docs)
    assert any("test DOCX file" in doc for doc in docs)

@pytest.mark.parametrize("sample_image_file", ["png", "jpg", "jpeg"], indirect=True)
def test_image_loader(sample_image_file):
    file, text = sample_image_file
    try:
        import pytesseract
        from pytesseract import TesseractNotFoundError
        # Check if tesseract is installed
        try:
            pytesseract.get_tesseract_version()
        except (TesseractNotFoundError, OSError):
            pytest.skip("Tesseract is not installed")
    except ImportError:
        pytest.skip("pytesseract is required for image test")
    loader = ImageLoader(str(file))
    loader.load()
    docs = loader.documents
    # OCR is not perfect, so just check that the expected word is in the output
    assert any("Test" in doc for doc in docs) 